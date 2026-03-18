#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建Word文档
"""

import sys
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

def get_alignment(alignment_str):
    """将字符串对齐方式转换为WD_ALIGN_PARAGRAPH枚举"""
    alignment_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    return alignment_map.get(alignment_str, WD_ALIGN_PARAGRAPH.LEFT)

def add_table(doc, table_config):
    """
    添加表格到文档
    
    Args:
        doc: Document对象
        table_config: 表格配置字典
            - headers: 表头列表
            - data: 表格数据（列表的列表）
            - style: 表格样式（可选，默认'Light Grid Accent 1'）
            - autofit: 是否自动调整列宽（可选，默认True）
    """
    headers = table_config.get('headers', [])
    data = table_config.get('data', [])
    style = table_config.get('style', 'Light Grid Accent 1')
    autofit = table_config.get('autofit', True)
    
    if not headers and not data:
        return
    
    # 确定行数和列数
    rows = len(data) + (1 if headers else 0)
    cols = max(len(headers) if headers else 0, 
               max([len(row) for row in data] if data else [0]))
    
    if cols == 0:
        return
    
    # 创建表格
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 填充表头
    if headers:
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            if i < cols:
                cell = header_cells[i]
                cell.text = str(header)
                # 设置表头样式
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    # 填充数据行
    data_row_idx = 1 if headers else 0
    for row_idx, row_data in enumerate(data):
        row_cells = table.rows[data_row_idx + row_idx].cells
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < cols:
                row_cells[col_idx].text = str(cell_data)
    
    # 应用表格样式
    table.style = style
    
    # 自动调整列宽
    if autofit:
        for column in table.columns:
            for cell in column.cells:
                cell.width = Inches(1.5)

def add_image(doc, image_config):
    """
    添加图片到文档
    
    Args:
        doc: Document对象
        image_config: 图片配置字典
            - path: 图片文件路径（必需）
            - width: 图片宽度（英寸，可选）
            - caption: 图片说明（可选）
            - alignment: 对齐方式（'left', 'center', 'right'，默认'center'）
    """
    image_path = image_config.get('path')
    if not image_path or not os.path.exists(image_path):
        print(f"⚠️ 图片文件不存在: {image_path}")
        return
    
    width = image_config.get('width', Inches(5))
    caption = image_config.get('caption')
    alignment = image_config.get('alignment', 'center')
    
    # 添加图片
    try:
        run = doc.add_picture(image_path, width=width)
        
        # 设置对齐方式
        if alignment:
            paragraph = run.paragraph
            paragraph.alignment = get_alignment(alignment)
        
        # 添加图片说明
        if caption:
            caption_para = doc.add_paragraph()
            caption_run = caption_para.add_run(f"图：{caption}")
            caption_run.italic = True
            caption_para.alignment = get_alignment(alignment)
            
    except Exception as e:
        print(f"❌ 无法添加图片 {image_path}: {str(e)}")

def create_word_document(output_path, title, content_data, subtitle=None, sections=None):
    """
    创建Word文档
    
    Args:
        output_path: 输出文件路径
        title: 文档标题
        content_data: 内容数据（可以是字符串或字典）
        subtitle: 副标题（可选）
        sections: 章节列表（可选）
    """
    
    # 创建文档
    doc = Document()
    
    # 设置文档属性
    doc.core_properties.author = "Document Generator"
    doc.core_properties.title = title
    
    # 设置标题
    doc_title = doc.add_heading(title, 0)
    doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加副标题
    if subtitle:
        sub_para = doc.add_paragraph()
        sub_run = sub_para.add_run(subtitle)
        sub_run.italic = True
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加日期
    date_para = doc.add_paragraph()
    date_para.add_run(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()  # 空行
    
    # 处理内容
    if isinstance(content_data, dict) and 'sections' in content_data:
        # 按章节组织
        for section in content_data['sections']:
            if 'heading' in section:
                level = section.get('level', 1)
                doc.add_heading(section['heading'], level=level)
            
            if 'content' in section:
                for item in section['content']:
                    if isinstance(item, dict):
                        # 复杂内容（带标题、列表、表格、图片等）
                        if 'type' in item:
                            if item['type'] == 'heading':
                                doc.add_heading(item['text'], level=item.get('level', 2))
                            elif item['type'] == 'list':
                                list_type = item.get('list_type', 'bullet')
                                for list_item in item.get('items', []):
                                    if list_type == 'numbered':
                                        p = doc.add_paragraph(list_item, style='List Number')
                                    else:
                                        p = doc.add_paragraph(list_item, style='List Bullet')
                                    p.paragraph_format.left_indent = Inches(0.25)
                            elif item['type'] == 'paragraph':
                                p = doc.add_paragraph()
                                run = p.add_run(item['text'])
                                if 'bold' in item and item['bold']:
                                    run.bold = True
                                if 'italic' in item and item['italic']:
                                    run.italic = True
                                if 'underline' in item and item['underline']:
                                    run.underline = True
                                if 'font_size' in item:
                                    run.font.size = Pt(item['font_size'])
                                if 'font_color' in item:
                                    run.font.color.rgb = RGBColor.from_string(item['font_color'])
                                p.alignment = get_alignment(item.get('alignment', 'left'))
                            elif item['type'] == 'table':
                                add_table(doc, item)
                            elif item['type'] == 'image':
                                add_image(doc, item)
                    else:
                        # 简单文本
                        doc.add_paragraph(str(item))
    else:
        # 简单文本内容
        if isinstance(content_data, list):
            for item in content_data:
                doc.add_paragraph(str(item))
        else:
            # 按行分割
            lines = str(content_data).split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    doc.add_paragraph(line)
    
    doc.add_paragraph()  # 空行
    
    # 页脚
    footer = doc.add_paragraph()
    footer.add_run('🐧 Document Generator')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 保存文档
    doc.save(output_path)
    print(f"✅ Word文档已创建: {output_path}")
    return output_path

def main():
    import argparse
    
    parser = argparse.ArgumentParser(description='创建Word文档')
    parser.add_argument('output', help='输出文件路径 (.docx)')
    parser.add_argument('--title', required=True, help='文档标题')
    parser.add_argument('--subtitle', help='文档副标题')
    parser.add_argument('--content-file', help='包含内容的文本文件')
    parser.add_argument('--content', help='文档内容（文本或JSON）')
    
    args = parser.parse_args()
    
    # 获取内容
    if args.content_file:
        with open(args.content_file, 'r', encoding='utf-8') as f:
            content_data = f.read()
    elif args.content:
        content_data = args.content
    else:
        print("❌ 必须提供 --content 或 --content-file")
        sys.exit(1)
    
    # 尝试解析JSON
    if content_data.strip().startswith('{'):
        try:
            import json
            content_data = json.loads(content_data)
        except:
            pass  # 保持原样
    
    # 创建文档
    create_word_document(
        output_path=args.output,
        title=args.title,
        content_data=content_data,
        subtitle=args.subtitle
    )

if __name__ == '__main__':
    main()